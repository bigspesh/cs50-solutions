allvalues = (1,2,3,4,3,5,6,4)
newvalues = []
for number in allvalues:
    if number not in newvalues:
       newvalues.append(number)
noduplicates = newvalues
print(noduplicates)
    

allvalues = [1,2,3,4,3,5,6,4]
newvalues = []
for number in allvalues:
    if allvalues.count(number) > 1:
        allvalues.remove(number)
newvalues1 = allvalues
print(newvalues1)

my_tuple = (1,2,3,4,5)
print(list(my_tuple))

s = "SIMPLILEARN"
for i in s:
    print(i, end="")

list2 = ['ctrl', 'John', "David"]
for i in range(len(list2)):
    print(f"Hello:{list2[i]}")


a = "aaa,bbb,ccc"
b = Counter(a)
print(list(a))


cities_in_F = {'New York': 32, 'Boston': 55, 'Denver': 65, 'Los_Angeles': 34}
a = Counter(cities_in_F)
print(a)


user_input = (input("Enter Change:"))
print("HELLO" + user_input)


car = 4.2
print("this is %d" %car )

text = input()
for i in text.split():
    print(i)

rom classautomation import get_validated_time
from docx import Document

document = Document()


def study_count():
    while True:
        try:
            class_count = int(input('Enter number of study periods in a day (3 or 4): '))
            if class_count == 4:
                table = document.add_table(5, 4)
                return table, class_count
            elif class_count == 3:
                table = document.add_table(5, 3)
                return table, class_count
            else:
                print('Study period should be 3 or 4.')
        except ValueError:
            print('Invalid input. Enter 3 or 4.')

# to unpack because calling the study_count() dunction alone returns a tuple in this form
# (<docx.table.Table object at 0x0000024CCC221AE0>, 4) and so table = <docx.table.Table object at 0x0000024CCC221AE0>
# which will then be used by then be used by the document app 
# and class_count = 4
# the essence of this is to assign values to our variables so you can reuse them anywhere within the code space
# eg print(f" the class count is : {class_count}")
# >> the class count is 4

table, class_count = study_count()



class_times = [
    'Enter the time interval for the first study period (e.g., 00:00AM/PM - 00:00AM/PM):',
    'Enter the time interval for the second study period (e.g., 00:00AM/PM - 00:00AM/PM):',
    'Enter the time interval for the third study period (e.g., 00:00AM/PM - 00:00AM/PM):',
    'Enter the time interval for the fourth study period (e.g., 00:00AM/PM - 00:00AM/PM):',
]

# To store the time intervals that have been used
used_intervals = set()

# Populate the first row with time intervals
# also while there is no unpacking going on here, we are using the each entry in the used intervals to set the text in the cell
# which is happening at runtime

for col in range(len(class_times)):
    if col < class_count:
        table.cell(0, col).text = get_validated_time(class_times[col], used_intervals)

# Prompt for the number of courses
def course_count_prompt():
    while True:
        try:
            course_count = int(input("Enter the number of courses you want to study (5-12): "))
            if 5 <= course_count <= 12:
                return course_count
            else:
                print("Course count should be between 5 and 12.")
        except ValueError:
            print("Invalid input. Enter a number between 5 and 12.")

# in this case an unpacking is also going on here, but this time the course_count_prompt() function is not 
# returning a tuple but rather returning the user_input which is course_count, then after you assign it to the variable course_count again
# like we did earlier in the table, class_count = study_count(), so we can reuse the value of course_count outside of this function
# eg print(f"the number of courses you'll study is {course_count})
# >> the number of courses you'll study is 12, that's if the user enter 12
# NOTE: if you dont assign the the course_count_prompt() function to the course_count variable, you ca use the course_count value outside of
# that fucnction, because he sole purpose of the function is to take the user_input and then return he value which you can then store inside
# of a variable 
course_count = course_count_prompt()

courses = {f'course{i}': input(f'Course{i}: ') for i in range(1, course_count + 1)}

# Populate the table with course names

# a funny code i wrot, you get 0 irespective of 
# the floating point number you enter 

def get():
    user_input = float(input())
    round = user_input // 1
    solve = user_input - round
    mult = solve * 100
    denom = [25, 9.999999999999972, 5.000000000000057, 0.9999999999999858]
    while mult >= denom[0]:
        mult -= denom[0] # decremnting 
    if mult == 0.0:
        print(mult)
    elif mult != 0.0:
        mult = mult - mult
    while mult >= denom[1]:
        mult = mult - denom[1]
        if mult == 0.0:
            print(mult)
        elif mult != 0.0:
            mult = mult - mult
    while mult >= denom[2]:
        mult = mult - denom[2]
    if mult == 0.0:
        print(mult)
    elif mult != 0.0:
        mult = mult - mult         
    while mult >= denom[3]:
        mult = mult - denom[3]
    if mult == 0.0:
        print(mult)
    elif mult != 0.0:
        mult = mult - mult       
    return mult , denom

mult , denom = get()
